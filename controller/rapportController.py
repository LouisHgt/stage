
import os
import docx
import subprocess
from docx.shared import Pt, RGBColor, Inches # Pour les unités et couleurs
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT # Pour l'alignement (si besoin)
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

class rapportController():
    def __init__(self, config_model_inst, couche_model_inst):
        self.configModel = config_model_inst
        self.coucheModel = couche_model_inst
        
    def setFormView(self, formView):
        self.formView = formView
        
    def setDialog(self, dialog):
        self.dialog = dialog
        
    def buildRapport(self):
        # On importe docx dans la methode pour eviter les conflits avec le garbage collector
        
        """
            Recupere les données d'entrée de la table 'site retenu' et crée un docx avec
        """
        
        # Recuparation de l'emplacement du rapport et de son nom
        emplacement_rapport = self.configModel.getFromConfig("emplacement_rapport")
        nom_rapport = self.configModel.getFromConfig("nom_rapport")
        
        rapport_path = os.path.join(os.path.dirname(__file__), '..', emplacement_rapport, nom_rapport) + ".docx"

        self.rapport = docx.Document()

        
        # Recuperation de la couche site_retenu
        emplacement_couche_site_retenu = self.configModel.getFromConfig('emplacement_couche_site_retenu')
        nom_couche_site_retenu = self.configModel.getFromConfig('nom_couche_site_retenu')
        path_site_retenu = os.path.join(os.path.dirname(__file__), '..', emplacement_couche_site_retenu, nom_couche_site_retenu) + ".shp"
        couche = self.coucheModel.getCoucheFromFile(path_site_retenu, "site_retenu")

        
        
        self.niveau = self.coucheModel.getNbrAttributsCouche(couche)
        self.list = [] # Liste dans laquelle on stocke les elements servants au filtre
        
        self.buildDocxRecursive(couche, self.coucheModel.getFilteredNiveau(couche))
        
        self.rapport.save(rapport_path)

        del self.rapport
        print("rapport ecrit")

        
        
        
        
    def buildDocxRecursive(self, couche,  liste_elements):
        """Fonction recursive qui parcours pour un elt d'un niveau les elt du niveau +1"""
        
        current_nv = len(self.list)
        # Condition d'arret
        if current_nv >= self.niveau - 1:
            for elt in liste_elements:
                pass
                p = self.rapport.add_paragraph(elt)
                self.apply_style_to_paragraph(p, current_nv)
            return
        
        for elt in liste_elements:
            

            self.list.append(elt)
            
            
            p = self.rapport.add_paragraph(elt)
            
            
            self.apply_style_to_paragraph(p, current_nv)
            
            
            # On filtre
            nv_liste = self.coucheModel.getFilteredNiveau(couche, self.list)
            # On rappelle avec la nouvelle liste
            self.buildDocxRecursive(couche, nv_liste)
            self.list.pop()
            

    def apply_style_to_paragraph(self, paragraph, level):
        para_format = paragraph.paragraph_format
        font = paragraph.runs[0].font if paragraph.runs else None # Assurer qu'il y a des runs

        # Police par défaut si non spécifiée autrement
        if font:
            font.name = 'Calibri'
            # Pour assurer la compatibilité avec les caractères complexes/asiatiques si besoin
            r = paragraph.runs[0]._r
            r.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

        # Reset format général
        para_format.left_indent = Inches(0)
        para_format.right_indent = Inches(0)
        para_format.first_line_indent = Inches(0)
        para_format.keep_together = False
        para_format.keep_with_next = False
        para_format.page_break_before = False
        para_format.widow_control = True # Empêche ligne seule en début/fin de page
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        if font:
            font.bold = False
            font.italic = False
            font.underline = False
            font.color.rgb = RGBColor(0, 0, 0) # Noir par défaut

        # --- Styles par niveau ---
        if level == 0:  # Niveau Bassin Versant (Titre Principal)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            if font:
                font.name = 'Arial'
                font.size = Pt(18)
                font.bold = True
                font.color.rgb = RGBColor(0x1F, 0x4E, 0x78) # Bleu foncé
                r = paragraph.runs[0]._r
                r.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
            para_format.space_before = Pt(24)
            para_format.space_after = Pt(12)
            para_format.keep_with_next = True
            para_format.page_break_before = True # Nouvelle page avant chaque grand titre

            # Ajouter une bordure inférieure
            p_borders = OxmlElement('w:pBdr')
            bottom_border = OxmlElement('w:bottom')
            bottom_border.set(qn('w:val'), 'single')
            bottom_border.set(qn('w:sz'), '6') # Taille en 1/8e de point (6 -> 0.75pt)
            bottom_border.set(qn('w:space'), '1')
            bottom_border.set(qn('w:color'), '1F4E78') # Couleur hexadécimale
            p_borders.append(bottom_border)
            paragraph._p.get_or_add_pPr().append(p_borders)


        elif level == 1: # Niveau Commune (Sous-titre 1)
            if font:
                font.name = 'Arial'
                font.size = Pt(14)
                font.bold = True
                font.color.rgb = RGBColor(0x2E, 0x74, 0xB5) # Bleu moyen
                r = paragraph.runs[0]._r
                r.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
            para_format.space_before = Pt(18)
            para_format.space_after = Pt(6)
            para_format.keep_with_next = True


        elif level == 2: # Niveau Type de site (Sous-titre 2)
            if font:
                font.name = 'Calibri'
                font.size = Pt(12)
                font.bold = True
                font.italic = False # Peut-être mettre en italique?
                font.color.rgb = RGBColor(0x59, 0x59, 0x59) # Gris foncé
                r = paragraph.runs[0]._r
                r.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
            para_format.left_indent = Inches(0.25)
            para_format.space_before = Pt(12)
            para_format.space_after = Pt(4)
            para_format.keep_with_next = True


        elif level == 3: # Niveau Site (Élément de liste)
             # Utiliser le style Liste à Puces s'il existe, sinon simuler
            try:
                 paragraph.style = 'List Bullet'
                 # Les styles peuvent définir retraits et espacements, mais on peut surcharger
                 para_format.space_before = Pt(2)
                 para_format.space_after = Pt(2)
                 if font:
                      font.size = Pt(11) # Ajuster la taille si besoin
                      font.color.rgb = RGBColor(0, 0, 0) # Assurer la couleur noire

            except KeyError:
                 # Simuler une liste si le style n'existe pas
                 if font:
                      font.name = 'Calibri'
                      font.size = Pt(11)
                      font.bold = False
                      font.italic = False
                      font.color.rgb = RGBColor(0, 0, 0)
                      r = paragraph.runs[0]._r
                      r.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

                 # Ajouter une puce manuellement (simple tiret ici)
                 if not paragraph.text.startswith("- "):
                       paragraph.text = "- " + paragraph.text

                 # Définir les retraits pour simuler une liste
                 para_format.left_indent = Inches(0.5) # Retrait global
                 para_format.first_line_indent = Inches(-0.25) # Retrait négatif pour la puce (hanging indent)
                 para_format.space_before = Pt(2)
                 para_format.space_after = Pt(4)
                 para_format.keep_with_next = False

        else: # Style par défaut pour niveaux non gérés ou texte normal
            if font:
                font.name = 'Calibri'
                font.size = Pt(11)
                r = paragraph.runs[0]._r
                r.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
            para_format.space_before = Pt(0)
            para_format.space_after = Pt(6)


    def convertToPdf(self):
        
        try:
            emplacement_rapport = self.configModel.getFromConfig("emplacement_rapport")
            nom_rapport = self.configModel.getFromConfig("nom_rapport")
            
            
            fichier_docx_entree = os.path.join(os.path.dirname(__file__), '..', emplacement_rapport) + nom_rapport + ".docx"
            fichier_pdf_sortie = os.path.join(os.path.dirname(__file__), '..', emplacement_rapport) + nom_rapport + '.pdf'
            dossier_pdf_sortie = os.path.join(os.path.dirname(__file__), '..', emplacement_rapport)
        
            lo_path = os.path.dirname("C:\Program Files\LibreOffice\program\soffice.exe")

            
            # Création de la commande
            command = [
                lo_path,
                '--headless', # Ne pas lancer l'interface graphique
                '--convert-to', 'pdf',
                '--outdir', dossier_pdf_sortie,
                fichier_docx_entree
            ]
        
            result = subprocess.run(
                command,
                capture_output=True,
                text=True,
                check=False,
                timeout=20
            )
            
            if result.returncode != 0:
                print("erreur lors de la conversion du docx en pdf pendant la commande lo")
        except Exception as e:
            print("Erreur lors de la conversion du docx en pdf")
            print(e)
            raise
        
    def handleFormTaskFinished(self, success):
        """
            Quand formTask est fini
        """
        
        self.buildRapport()
        self.dialog.accept()