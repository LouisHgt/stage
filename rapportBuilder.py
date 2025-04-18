
import os
import docx
from docx.shared import Pt, RGBColor, Cm # Pour les unités et couleurs
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT # Pour l'alignement (si besoin)
from .configManager import configManager
from docx2pdf import convert

class rapportBuilder():
    def __init__(self, coucheManager):
        self.configManager = configManager()
        self.coucheManager = coucheManager

    def buildRapport(self, fileType1):
        # On importe docx dans la methode pour eviter les conflits avec le garbage collector
        
        """Recupere les données d'entrée de la table 'site retenu' et crée un docx avec

        Args:
            Prend entre 1 et deux types de documents de sortie
        """
        
        # Recuparation de l'emplacement du rapport et de son nom
        emplacement_rapport = self.configManager.getFromConfig("emplacement_rapport")[0]
        nom_rapport = self.configManager.getFromConfig("nom_rapport")[0]
        
        rapport_path = os.path.join(os.path.dirname(__file__), emplacement_rapport, nom_rapport) + fileType1

        self.rapport = docx.Document()
        
        # Recuperation de la couche site_retenu
        emplacement_couche_site_retenu = self.configManager.getFromConfig('emplacement_couche_site_retenu')[0]
        nom_couche_site_retenu = self.configManager.getFromConfig('nom_couche_site_retenu')[0]
        path_site_retenu = os.path.join(os.path.dirname(__file__), emplacement_couche_site_retenu, nom_couche_site_retenu) + ".shp"
        couche = self.coucheManager.getCoucheFromFile(path_site_retenu, "site_retenu")

        
        
        self.niveau = self.coucheManager.getNbrAttributsCouche(couche)
        self.list = [] # Liste dans laquelle on stocke les elements servants au filtre
        self.buildDocxRecursive(couche, self.coucheManager.getFilteredNiveau(couche))
        
        
        self.rapport.save(rapport_path)
        del self.rapport
        print("rapport ecrit")

        
        
        
        
    def buildDocxRecursive(self, couche,  liste_elements):
        """Fonction recursive qui parcours pour un elt d'un niveau les elt du niveau +1"""
        
        current_nv = len(self.list)
        # Condition d'arret
        if current_nv >= self.niveau - 1:
            for elt in liste_elements:
                p = self.rapport.add_heading(elt, 1)
                self.apply_style_to_paragraph(p, current_nv)
            return
        
        for elt in liste_elements:
            

            self.list.append(elt)
            
            
            p = self.rapport.add_paragraph(elt)
            self.apply_style_to_paragraph(p, current_nv)
            
            
            
            # On filtre
            nv_liste = self.coucheManager.getFilteredNiveau(couche, self.list)
            # On rappelle avec la nouvelle liste
            self.buildDocxRecursive(couche, nv_liste)
            self.list.pop()
            

    def apply_style_to_paragraph(self, paragraph, level):
        """
        Applique un style et/ou un formatage direct au paragraphe
        selon son niveau hiérarchique (0=Bassin, 1=Commune, 2=Type, 3=Site).
        """
        # --- Références rapides ---
        para_format = paragraph.paragraph_format
        font = None
        if paragraph.runs:
            font = paragraph.runs[0].font

        # --- Application des styles/formats par niveau ---
        if level == 0:  # Niveau Bassin Versant (Titre Principal)
            try:
                paragraph.style = 'Heading 1' # Style Word intégré
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                font.color.rgb = RGBColor(255, 0, 0)
            except KeyError:
                print("Avertissement: Style 'Heading 1' non trouvé, utilisation de formatage direct.")
                if font:
                    font.name = 'Calibri' # Ou autre police de titre
                    font.size = Pt(16)
                    font.bold = True
                    font.color.rgb = RGBColor(0x1F, 0x4E, 0x78) # Exemple: Bleu foncé
            # Espacement
            para_format.space_before = Pt(18)
            para_format.space_after = Pt(6)
            para_format.keep_with_next = True # Évite coupure avant le contenu

        elif level == 1: # Niveau Commune (Sous-titre 1)
            try:
                paragraph.style = 'Heading 2'
                font.color.rgb = RGBColor(6, 0, 157)
            except KeyError:
                print("Avertissement: Style 'Heading 2' non trouvé.")
                if font:
                    font.name = 'Calibri'
                    font.size = Pt(14)
                    font.bold = True
            # Espacement
            para_format.space_before = Pt(12)
            para_format.space_after = Pt(4)
            para_format.keep_with_next = True

        elif level == 2: # Niveau Type de site (Sous-titre 2)
            try:
                paragraph.style = 'Heading 3'
                para_format.first_line_indent = Cm(1.27)
            except KeyError:
                print("Avertissement: Style 'Heading 3' non trouvé.")
                if font:
                    font.name = 'Calibri'
                    font.size = Pt(12)
                    font.bold = True
                    # font.italic = True # Optionnel
            # Espacement
            para_format.space_before = Pt(10)
            para_format.space_after = Pt(2)
            para_format.keep_with_next = True

        elif level == 3: # Niveau Site (Élément de liste)
            try:
                paragraph.style = 'List Bullet'
                para_format.first_line_indent = Cm(1.27)
            except KeyError:
                print("Avertissement: Style 'List Bullet' non trouvé, utilisation formatage simple.")
                if font:
                    font.name = 'Calibri'
                    font.size = Pt(11)
                    font.bold = False
                # Ajouter un retrait manuel si le style n'existe pas
                para_format.left_indent = Cm(0.75)
            # Espacement entre les elements de la liste
            para_format.space_before = Pt(0)
            para_format.space_after = Pt(4)
            para_format.keep_with_next = False # Permet les coupures entre les items

        else: # Style par défaut pour niveaux non gérés
            try:
                paragraph.style = 'Normal'
            except KeyError:
                 print("Avertissement: Style 'Normal' non trouvé.")
                 # Appliquer un formatage de base
                 if font:
                      font.name = 'Calibri'
                      font.size = Pt(11)
            para_format.space_before = Pt(0)
            para_format.space_after = Pt(6)


    def convertToPdf(self):
        fichier_docx_entree = "votre_document.docx"
        fichier_pdf_sortie = "votre_document.pdf" 
        #Todo : conversion du fichier docx créé en pdf